import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentExtractionError(Exception):
    """Raised when document text extraction fails."""


def clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", "")

    text = re.sub(
        r"[ \t]+",
        " ",
        text,
    )

    text = re.sub(
        r"\n\s*\n\s*\n+",
        "\n\n",
        text,
    )

    return text.strip()


def extract_pdf_text(storage_path: Path) -> str:
    try:
        reader = PdfReader(str(storage_path))
    except Exception as exc:
        raise DocumentExtractionError(
            "The PDF file could not be opened."
        ) from exc

    if reader.is_encrypted:
        try:
            result = reader.decrypt("")
        except Exception as exc:
            raise DocumentExtractionError(
                "Encrypted PDF files are not supported."
            ) from exc

        if result == 0:
            raise DocumentExtractionError(
                "Password-protected PDF files are not supported."
            )

    page_texts: list[str] = []

    try:
        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):
            page_text = page.extract_text() or ""

            if page_text.strip():
                page_texts.append(
                    f"--- Page {page_number} ---\n"
                    f"{page_text}"
                )

    except Exception as exc:
        raise DocumentExtractionError(
            "Text could not be extracted from the PDF."
        ) from exc

    return clean_extracted_text(
        "\n\n".join(page_texts)
    )


def extract_docx_text(storage_path: Path) -> str:
    try:
        document = DocxDocument(
            str(storage_path)
        )
    except Exception as exc:
        raise DocumentExtractionError(
            "The DOCX file could not be opened."
        ) from exc

    content_parts: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            content_parts.append(
                paragraph_text
            )

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )

            if row_text:
                content_parts.append(
                    row_text
                )

    return clean_extracted_text(
        "\n\n".join(content_parts)
    )


def extract_document_text(
    storage_path: str | Path,
    content_type: str,
    original_filename: str,
) -> str:
    path = Path(storage_path).resolve()

    if not path.exists():
        raise DocumentExtractionError(
            "The stored document file does not exist."
        )

    if not path.is_file():
        raise DocumentExtractionError(
            "The stored document path is not a file."
        )

    filename_extension = (
        Path(original_filename)
        .suffix
        .lower()
    )

    if (
        content_type == "application/pdf"
        or filename_extension == ".pdf"
    ):
        extracted_text = extract_pdf_text(path)

    elif (
        content_type
        == (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
        or filename_extension == ".docx"
    ):
        extracted_text = extract_docx_text(path)

    else:
        raise DocumentExtractionError(
            "This document type is not supported."
        )

    if not extracted_text:
        raise DocumentExtractionError(
            "No readable text was found. The document may "
            "contain only scanned images."
        )

    return extracted_text